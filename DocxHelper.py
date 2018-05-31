#!/usr/bin/python3
# -*- coding: UTF-8 -*-


from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor

import docx
from Config import log

class DocxHelper():
    def __init__(self, file_name):
        self.doc = Document()
        self.file_name = file_name

    def add_hyperlink(self, paragraph, url, text):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        # This gets access to the document.xml.rels file and gets a new relation id value
        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create the w:hyperlink tag and add needed values
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

        # Create a w:r element
        new_run = docx.oxml.shared.OxmlElement('w:r')

        # Create a new w:rPr element
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        # Add color if it is given

        c = docx.oxml.shared.OxmlElement('w:rFonts')
        c.set(docx.oxml.shared.qn('w:hAnsi'), '仿宋')
        c.set(docx.oxml.shared.qn('w:eastAsia'), '仿宋')
        c.set(docx.oxml.shared.qn('w:ascii'), '仿宋')
        rPr.append(c)
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), '548DD4')
        c.set(docx.oxml.shared.qn('w:themeTint'), '99')
        c.set(docx.oxml.shared.qn('w:themeColor'), 'text2')
        rPr.append(c)
        c = docx.oxml.shared.OxmlElement('w:sz')
        c.set(docx.oxml.shared.qn('w:val'), '44')
        rPr.append(c)
        c = docx.oxml.shared.OxmlElement('w:szCs')
        c.set(docx.oxml.shared.qn('w:val'), '44')
        rPr.append(c)

        # Join all the xml elements together add add the required text to the w:r element
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

        return hyperlink

    def save(self):
        self.doc.save(self.file_name)

    def add_paragragh(self, text, color, images, bold, links=None):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = u'仿宋'
        run.font.size = Pt(22)
        run.font.color.rgb = color
        run.font.bold = bold
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
        if images:
            for image in images:
                try:
                    self.doc.add_picture(image)
                except Exception as e:
                    log('add image {} failed. {}'.format(image, e))
        if links:
            for link in links:
                self.add_hyperlink(p, link[1], '按住Ctrl点击查看链接:' + link[0])
                # p.add_hyperlink( link[1], '按住Ctrl点击查看链接:' +link[0],None,False)
                # hyperlink.font = run.font
        return p

    def add_QA(self, question, answer, create_time, images_q, images_a, links):
        self.add_paragragh("【{} {}】问题:{}".format(create_time[:10], create_time[11:19], question), RGBColor(0, 0, 0),
                           images_q, False, links)
        self.add_paragragh("【老齐回答】:{}\n".format(answer), RGBColor(139, 58, 58), images_a, True, [])

    def add_Talk(self, text, create_time, images, links, file_links):
        p = self.add_paragragh("【{} {}】{}".format(create_time[:10], create_time[11:19], text), RGBColor(0, 0, 0),
                               images, False, links)
        for link in file_links:
            self.add_hyperlink(p, link[1], '按住Ctrl点击下载:' + link[0])


if __name__ == '__main__':
    d = DocxHelper('test.docx')
    d.add_QA('标题', '回答', '2018-12-12T12:23:23', [], [], [('abcdddd', 'http://163.com')])
    d.add_Talk('内容', '2018-12-12T12:23:23', [], [], [('文件下载', 'http://163.com')])
    d.save()
    # import re
    # import urllib.request
    #
    # text = r'<e type="web" href="https%3A%2F%2Fmp.weixin.qq.com%2Fs%2FK4jIFpdrdECvROQASeZpvA" title="%E7%8E%AF%E7%90%83%E6%97%B6%E5%B1%80" cache="http%3A%2F%2Fcache.zsxq.com%2F201805%2F371f7956cd72ae85f0c3a465bab7a226054b3d74eca61b1c440cbbfd9e6bf368"/>'
    # result = re.findall('href="(.*?)[\s]*title="(.*?)"', urllib.request.unquote(text))  # title="(.*?)"
    # print(result[0][1])

    #
    # def add_hyperlink(paragraph, url, text, color, underline):
    #        """
    #        A function that places a hyperlink within a paragraph object.
    #
    #        :param paragraph: The paragraph we are adding the hyperlink to.
    #        :param url: A string containing the required url
    #        :param text: The text displayed for the url
    #        :return: The hyperlink object
    #        """
    #
    #        # This gets access to the document.xml.rels file and gets a new relation id value
    #        part = paragraph.part
    #        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    #
    #        # Create the w:hyperlink tag and add needed values
    #        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    #        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    #
    #        # Create a w:r element
    #        new_run = docx.oxml.shared.OxmlElement('w:r')
    #
    #        # Create a new w:rPr element
    #        rPr = docx.oxml.shared.OxmlElement('w:rPr')
    #        # Add color if it is given
    #
    #
    #        c = docx.oxml.shared.OxmlElement('w:rFonts')
    #        c.set(docx.oxml.shared.qn('w:hAnsi'), '仿宋')
    #        c.set(docx.oxml.shared.qn('w:eastAsia'), '仿宋')
    #        c.set(docx.oxml.shared.qn('w:ascii'), '仿宋')
    #        rPr.append(c)
    #        c = docx.oxml.shared.OxmlElement('w:color')
    #        c.set(docx.oxml.shared.qn('w:val'), '548DD4')
    #        c.set(docx.oxml.shared.qn('w:themeTint'), '99')
    #        c.set(docx.oxml.shared.qn('w:themeColor'), 'text2')
    #        rPr.append(c)
    #        c = docx.oxml.shared.OxmlElement('w:sz')
    #        c.set(docx.oxml.shared.qn('w:val'), '44')
    #        rPr.append(c)
    #        c = docx.oxml.shared.OxmlElement('w:szCs')
    #        c.set(docx.oxml.shared.qn('w:val'), '44')
    #        rPr.append(c)
    #
    #
    #        # Join all the xml elements together add add the required text to the w:r element
    #        new_run.append(rPr)
    #        new_run.text = text
    #        hyperlink.append(new_run)
    #        paragraph._p.append(hyperlink)
    #
    #        return hyperlink
